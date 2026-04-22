from __future__ import annotations

from html import escape
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


OUT = Path(r"E:\T\hardsecnet\Reports\blackbook_recreated_local_only")
OUT.mkdir(parents=True, exist_ok=True)

PROJECT = "HardSecNet"
SUBTITLE = "A Local-First Cross-Platform Security Hardening Studio"
YEAR = "2025-26"
COLLEGE = "Walchand Institute of Technology, Solapur"
UNIVERSITY = "Punyashlok Ahilyadevi Holkar Solapur University, Solapur"
SUPERVISOR = "Dr. Mrs. R. J. Shelke"
HOD = "Dr. S. R. Gengaje"
PRINCIPAL = "Dr. V. A. Athavale"

STUDENTS = [
    ("Rahul Ravindra Nawale", "04", "2203111059"),
    ("Tushar Jitendra Satpute", "05", "2203111006"),
    ("Siddharth Bharat Magdum", "06", "2203111023"),
]

INDEX_ROWS = [
    ["1", "Introduction & Phase-I Summary", "Domain background, problem statement, objectives, phase-I summary", "4", "7"],
    ["2", "System Architecture & Development Approach", "Architecture, components, software flow, tools", "1", "11"],
    ["3", "Implementation & System Integration", "Implementation details, integration, developed UI", "1,3", "19"],
    ["4", "Testing, Experimentation & Data Acquisition", "Testing setup, data acquisition, observations", "2", "24"],
    ["5", "Performance Evaluation & Validation", "Results, analysis, validation outcomes", "2", "31"],
    ["6", "Environmental & Safety Considerations and SDG Alignment", "Environmental impact, safety, risk mitigation, SDGs", "2", "37"],
    ["7", "Conclusion & Future Scope", "Summary, achievements, limitations, future scope", "3,4", "40"],
    ["8", "References", "IEEE-style sources", "-", "43"],
]

TOOLS_TABLE = [
    ["Layer", "Technology", "Purpose"],
    ["Desktop UI", "PySide6", "Local dashboard and operator workflow"],
    ["Language", "Python 3.12", "Application logic and tests"],
    ["Database", "SQLite", "Local persistence"],
    ["Reports", "JSON, HTML, PDF", "Evidence export"],
    ["Benchmarks", "CIS Windows 11, CIS Ubuntu 24.04", "Authoritative control source"],
    ["AI", "Ollama-oriented local settings", "Risk explanation path"],
    ["Testing", "pytest", "Automated validation"],
]

TEST_TABLE = [
    ["Test Case", "Purpose", "Status"],
    ["TC-01", "Bootstrap local repository and dashboard snapshot", "Pass"],
    ["TC-02", "Run profile and create report artifacts", "Pass"],
    ["TC-03", "Import benchmark and create profile/document", "Pass"],
    ["TC-04", "Build PySide window and refresh all pages", "Pass"],
    ["TC-05", "Verify Dashboard is first local page", "Pass"],
]

VALIDATION_TABLE = [
    ["Validation Item", "Observed Result"],
    ["Compile source and tests", "Passed"],
    ["pytest test suite", "10 passed"],
    ["Navigation", "7 local pages with Dashboard first"],
    ["Committed Windows controls", "477"],
    ["Committed Ubuntu controls", "312"],
    ["Committed script candidates", "789"],
    ["Removed fleet implementation references", "No active source/test matches"],
]

APPENDICES = [
    (
        "APPENDIX A: CURRENT REPOSITORY EVIDENCE",
        [
            "The final implementation is represented by the hardsecnet-pyside repository. The important runtime modules are app.py, services.py, persistence.py, benchmark.py, agents.py, config.py, models.py, and the PySide6 UI pages under src/hardsecnet_pyside/ui.",
            "The remote control-plane, child-device agent, shared remote contracts, and web fleet dashboard were removed from active implementation scope. The remaining project can be executed and validated on one local system.",
            "The repository includes a README, First Run guide, Architecture guide, CIS Benchmark Engine guide, validation notes, BMAD PRD, BMAD architecture, BMAD development plan, and BMAD checkpoint artifacts.",
        ],
    ),
    (
        "APPENDIX B: BENCHMARK DATASET SUMMARY",
        [
            "The committed Windows benchmark export is named cis-microsoft-windows-11-stand-alone-benchmark-v4.0.0. It contains 477 structured controls and 477 script candidates.",
            "The committed Ubuntu benchmark export is named cis-ubuntu-linux-24.04-lts-benchmark-v1.0.0. It contains 312 structured controls and 312 script candidates.",
            "Each benchmark export includes benchmark_document.json, benchmark_items.json, scripts, and README.md. These exports allow the desktop app to load benchmark data without requiring the original PDF files at runtime.",
        ],
    ),
    (
        "APPENDIX C: DESKTOP UI SURFACE",
        [
            "The final PySide6 application contains seven local pages: Dashboard, Hardening, Network, AI Advisor, Reports, Benchmarks, and Settings.",
            "Dashboard summarizes current-device posture. Hardening handles profile selection and local run creation. Network shows posture checks. AI Advisor shows local risk/remediation explanation records. Reports shows exported artifacts. Benchmarks supports benchmark inspection and import. Settings shows runtime paths and AI mode.",
            "The interface uses metric cards, framed panels, readable tables, and a left navigation rail so that the software appears as a complete operator tool rather than as a raw database viewer.",
        ],
    ),
    (
        "APPENDIX D: VALIDATION COMMANDS",
        [
            "The implementation was validated with python -m compileall src tests.",
            "The automated test suite was validated with python -m pytest -q tests, producing 10 passed tests.",
            "A source scan over active src and tests files confirmed that no remote fleet implementation references remain after pruning.",
        ],
    ),
    (
        "APPENDIX E: SCOPE CORRECTION RECORD",
        [
            "The project was corrected from a broader fleet-control direction to a local-first CIS hardening studio. This correction matches the practical constraint that the project must be demonstrable without multiple PCs.",
            "The final scope includes CIS benchmark controls, script candidates, current-device dashboard, local profile execution, drift comparison, local AI explanation records, and reports.",
            "The final scope excludes parent admin PC control of child PCs, remote enrollment, heartbeat loops, remote jobs, campaigns, web fleet dashboards, and cloud AI dependency.",
        ],
    ),
]

CHAPTERS = [
    (
        "Chapter: 1",
        "INTRODUCTION & PROJECT PHASE-I SUMMARY",
        [
            (
                "1.1 Background of the Domain and Relevance",
                [
                    "System security hardening is the practice of reducing the attack surface of a computing system by disabling unsafe defaults, enforcing secure configurations, collecting audit evidence, and documenting remediation actions. Modern desktop and server operating systems contain thousands of configuration options across account policy, audit policy, network services, file permissions, kernel parameters, firewall settings, and application control.",
                    "The Center for Internet Security (CIS) publishes benchmark documents that are widely used as practical hardening baselines. CIS benchmarks describe expected configuration states, rationales, audit checks, remediation guidance, and the possible operational impact of each control. They are useful for both professional administrators and students because they connect security theory with concrete system settings.",
                    "HardSecNet applies this benchmark-driven approach to a local-first desktop application. Instead of building a fleet management platform that requires multiple physical machines, the final project focuses on a single operator device. The application imports and stores benchmark controls, maps them to reviewable script candidates, runs local profile-based hardening workflows, records evidence, compares before and after drift, explains risk using local AI-oriented records, and exports reports.",
                    "This scope is relevant to Electronics and Computer Engineering because it combines operating system configuration, scripting, local persistence, GUI development, evidence generation, and security documentation in one software project.",
                ],
            ),
            (
                "1.2 Problem Statement and Scope",
                [
                    "CIS benchmark documents are authoritative but large, dense, and difficult to operationalize manually. A student or administrator must read long PDF documents, identify relevant controls, translate them into PowerShell or shell logic, capture proof before and after execution, and prepare a readable report.",
                    "HardSecNet addresses the absence of a simple local tool that connects benchmark content, script candidates, local execution records, drift comparison, AI-assisted risk explanation, and report generation in a single workflow. The project is software-only and local-first. It does not depend on a remote server, child machines, device agents, or fleet campaigns.",
                    "The final project scope covers Windows 11 and Ubuntu 24.04 LTS benchmark content. It includes imported benchmark bundles, generated script candidates, local PySide6 dashboard pages, SQLite persistence, local report generation, deterministic AI explanation records, and validation tests.",
                ],
            ),
            (
                "1.3 Objectives of the Project",
                [
                    "Develop a CIS benchmark ingestion pipeline that extracts and normalizes controls from Windows 11 and Ubuntu benchmark material into structured local records.",
                    "Generate reviewable PowerShell and shell script candidates mapped to benchmark controls with traceable identifiers.",
                    "Implement a PySide6 desktop application that works without a backend server and opens with a current-device dashboard.",
                    "Support local profile execution, findings generation, evidence recording, approval records, and local report export.",
                    "Provide before and after drift comparison between local runs so that posture movement can be reviewed.",
                    "Provide local AI-oriented risk and remediation explanation records with Ollama-ready settings and a deterministic fallback.",
                ],
            ),
            (
                "1.4 Summary of Work Completed in Project Phase-I",
                [
                    "During Project Phase-I, the team identified CIS benchmark hardening as the main problem domain and studied Windows and Linux security benchmark structures. The early implementation explored benchmark ingestion, GUI-based control browsing, and evidence/report generation.",
                    "The initial repository also experimented with a broader fleet-control direction. During final scope correction, that direction was removed because the core academic objective could be demonstrated more reliably through a local-first security hardening studio.",
                    "The Phase-II implementation concentrates on the local desktop product: benchmark library, script candidates, current-device dashboard, local run records, drift comparison, local AI explanation, and reports.",
                ],
            ),
        ],
    ),
    (
        "Chapter: 2",
        "SYSTEM ARCHITECTURE & DEVELOPMENT APPROACH",
        [
            (
                "2.1 High-Level System Architecture of the Developed System",
                [
                    "HardSecNet PySide is structured as a single-process desktop application. The PySide6 user interface communicates with a controller facade, which delegates workflow operations to the service layer. The service layer uses a local SQLite repository, benchmark importer, local AI explanation engine, and report generator.",
                    "The current application exposes seven local pages: Dashboard, Hardening, Network, AI Advisor, Reports, Benchmarks, and Settings. The Dashboard is the first page and presents benchmark-control count, local run count, review items, drift deltas, AI task count, report count, latest findings, and drift movement.",
                ],
            ),
            (
                "2.2 Description of Software Modules and Components Used",
                [
                    "The controller module provides a stable API for UI pages and hides persistence details from the desktop interface.",
                    "The service module performs bootstrap, profile execution, findings generation, comparison generation, report creation, and AI task storage.",
                    "The persistence module stores local records in SQLite, including devices, profiles, benchmark documents, benchmark items, runs, findings, comparisons, reports, approvals, AI tasks, and settings.",
                    "The benchmark module imports documents, extracts control records, writes durable benchmark exports, and generates script candidates.",
                    "The agent module creates local explanation records for ingestion, reasoning, remediation planning, approval gating, and report writing.",
                ],
            ),
            (
                "2.3 Software Architecture, Algorithms, and Program Flow",
                [
                    "On startup, the application discovers runtime paths, creates local folders, initializes SQLite tables, loads built-in profiles, and refreshes committed benchmark export bundles.",
                    "For benchmark import, the importer reads the source file, creates a benchmark document, extracts control-like entries, builds audit/remediation metadata, writes generated scripts, and stores both document and item records locally.",
                    "For profile execution, the service selects benchmark items for the current device operating system, creates module results, synthesizes findings, compares against the previous local run, generates reports, and stores AI explanation records.",
                    "For drift comparison, the application compares previous and current findings by benchmark identifier and records whether controls improved, regressed, changed, or remained unchanged.",
                ],
            ),
            (
                "2.4 Software Tools and Platforms",
                [
                    "Python 3.12 is used as the primary programming language. PySide6 provides the desktop GUI. PyMuPDF supports PDF processing and report output. SQLite provides local persistence. Tests are written with pytest. The AI settings are local and Ollama-oriented.",
                ],
            ),
        ],
    ),
    (
        "Chapter: 3",
        "IMPLEMENTATION & SYSTEM INTEGRATION",
        [
            (
                "3.1 Software Implementation and Program Development",
                [
                    "The final implementation lives in the hardsecnet-pyside project. The desktop entry point creates a QApplication, constructs the MainWindow, and loads the HardSecNetController. The controller bootstraps the HardSecNetService and exposes methods used by UI pages.",
                    "The Dashboard page was added as the first screen to make the product feel like a complete hardening studio rather than a collection of raw tables. It summarizes the current device, benchmark controls, local runs, review items, drift deltas, AI tasks, and reports.",
                    "The Hardening page allows profile selection and local execution. The Network page summarizes network-related posture findings. The AI Advisor page displays local explanation records. The Reports page shows exported report artifacts. The Benchmarks page supports benchmark browsing and import. The Settings page displays runtime paths and local AI mode.",
                ],
            ),
            (
                "3.2 Module Integration",
                [
                    "Integration is performed through service-layer calls rather than direct database access from the UI. This keeps GUI code simpler and allows the repository schema to evolve without changing page logic.",
                    "Benchmark exports are committed under src/hardsecnet_pyside/data/benchmark_exports. The repository refreshes these bundles during bootstrap, so the application does not need the original CIS PDFs at runtime for the seeded benchmark content.",
                    "Reports are generated under runtime/reports and include JSON, HTML, and PDF artifacts. Run artifacts are stored under runtime/artifacts. Generated script candidates are stored under runtime/generated_scripts during imports and under committed benchmark export bundles for seeded content.",
                ],
            ),
            (
                "3.3 Developed System Screens and User Flow",
                [
                    "The final user flow begins on the Dashboard. The operator checks posture metrics, opens Benchmarks to inspect imported CIS controls, opens Hardening to run a selected profile, reviews findings, checks AI Advisor explanations, observes drift movement, and exports reports.",
                    "The UI uses a multi-accent dark desktop visual system with compact metric cards, framed panels, readable tables, and a left navigation rail. The color treatment separates benchmark metrics, review state, drift state, AI state, and report state without relying on a single-hue palette.",
                ],
            ),
        ],
    ),
    (
        "Chapter: 4",
        "TESTING, EXPERIMENTATION & DATA ACQUISITION",
        [
            (
                "4.1 Experimental Setup and Testing Procedures",
                [
                    "Testing was performed on a Windows development machine using the project virtual environment. The final verification commands were python -m compileall src tests and python -m pytest -q tests. The active test suite passed with 10 tests.",
                    "The UI smoke test creates a QApplication, builds the MainWindow, verifies that the Dashboard is the first navigation item, checks the expected page count, refreshes all pages, and verifies the status bar output. Controller tests verify bootstrap, profile execution, report creation, AI task creation, benchmark import, and report payload export.",
                ],
            ),
            (
                "4.2 Data Acquisition and Measurement Techniques",
                [
                    "Benchmark data is acquired from committed export bundles generated from CIS Windows 11 and CIS Ubuntu 24.04 LTS content. Each bundle contains a benchmark_document.json file, a benchmark_items.json file, scripts, and a README file.",
                    "The committed Windows benchmark bundle contains 477 controls and 477 script candidates. The committed Ubuntu benchmark bundle contains 312 controls and 312 script candidates. These records form the main dataset for local dashboard and report validation.",
                ],
            ),
            (
                "4.3 Experimental Observations and Test Cases",
                [
                    "The application successfully boots without a server, loads benchmark records, creates runtime folders, runs profiles, stores findings, writes reports, and displays the local dashboard. Source scans after scope pruning showed no active implementation references to removed remote fleet surfaces.",
                ],
            ),
        ],
    ),
    (
        "Chapter: 5",
        "PERFORMANCE EVALUATION & VALIDATION",
        [
            (
                "5.1 Results",
                [
                    "The final local-only project produced a working PySide6 desktop application with seven local pages and a corrected scope. The committed benchmark export data contains 789 total benchmark controls and 789 script candidates across Windows and Ubuntu bundles.",
                    "The app validates the main academic workflow: benchmark library to local profile run, local profile run to findings, findings to drift comparison, drift and findings to AI explanation records, and report generation to local artifacts.",
                ],
            ),
            (
                "5.2 Analysis",
                [
                    "The local-first design reduces deployment complexity because no server, agent, or web dashboard is required. It also makes the project easier to demonstrate in a classroom or viva environment because the complete workflow can be shown on one machine.",
                    "The main technical limitation is that local hardening execution is still scaffolded and review-gated. Running generated hardening scripts without validation could change live operating system policy, network behavior, or login settings. Therefore, the project records findings and script candidates but keeps production script execution as a controlled future enhancement.",
                ],
            ),
            (
                "5.3 Validation of Expected Outcomes",
                [
                    "Expected outcome one, local operation without a server, was satisfied. Expected outcome two, benchmark record storage and browsing, was satisfied through committed export bundles and UI pages. Expected outcome three, local profile execution and report generation, was satisfied through controller and service tests. Expected outcome four, drift comparison, was satisfied through comparison records between runs. Expected outcome five, AI explanation storage, was satisfied through local AI task records. Expected outcome six, fleet-scope removal, was satisfied by deleting the remote services and scanning active source/test files.",
                ],
            ),
        ],
    ),
    (
        "Chapter: 6",
        "ENVIRONMENTAL & SAFETY CONSIDERATIONS AND SDG ALIGNMENT",
        [
            (
                "6.1 Environmental Impact and Sustainability",
                [
                    "The project is software-only and runs locally on existing computing hardware. It does not require additional embedded devices, servers, routers, or endpoint machines for the final demonstration. This reduces hardware consumption and makes the project more sustainable for academic evaluation.",
                ],
            ),
            (
                "6.2 Energy Efficiency",
                [
                    "The local-first architecture avoids always-on server infrastructure. The application runs only when the operator needs to review benchmarks, execute a profile, or generate reports. SQLite storage and local files have minimal energy overhead compared with distributed services.",
                ],
            ),
            (
                "6.3 Safety Measures",
                [
                    "Security hardening can affect login behavior, service availability, network access, and audit logging. Therefore, HardSecNet treats generated scripts as reviewable candidates. The report and UI emphasize approval, rollback awareness, and evidence capture instead of blind execution.",
                ],
            ),
            (
                "6.4 Risk Mitigation",
                [
                    "The main risk is unsafe remediation on a live system. The mitigation is to keep script execution approval-gated, preserve rollback notes, record evidence, and classify script readiness before production use. Another risk is overclaiming AI output. The current system records deterministic local explanations and identifies live Ollama integration as future work.",
                ],
            ),
            (
                "6.5 Sustainable Development Goals Addressed",
                [
                    "The project supports SDG 4, Quality Education, by giving students a practical way to understand benchmark-based security hardening. It supports SDG 9, Industry, Innovation and Infrastructure, by improving local security tooling and auditability. It supports SDG 12, Responsible Consumption and Production, by avoiding unnecessary extra hardware for demonstration.",
                ],
            ),
        ],
    ),
    (
        "Chapter: 7",
        "CONCLUSION & FUTURE SCOPE",
        [
            (
                "7.1 Summary of Work Carried Out in Project Phase-II",
                [
                    "Project Phase-II produced a local-first PySide6 security hardening studio aligned with CIS benchmark workflows. The final implementation includes benchmark exports, script candidates, a local dashboard, profile execution scaffolding, findings, drift comparison, AI explanation records, reports, documentation, and tests.",
                    "The project was corrected from an overly broad fleet-control direction into a realistic software-only final product. This makes the final system easier to validate and defend because it does not require multiple endpoint machines.",
                ],
            ),
            (
                "7.2 Achievements and Limitations",
                [
                    "Achievements include a polished local desktop dashboard, durable benchmark exports, 789 committed benchmark controls, 789 script candidates, SQLite persistence, report generation, and a passing 10-test suite. The major limitation is that hardening execution and live Ollama model calls are not fully production-wired yet.",
                ],
            ),
            (
                "7.3 Future Scope",
                [
                    "Future work includes a validated local script runner with pre-check, apply, post-check, rollback, and evidence capture. Live Ollama integration can be added for richer risk explanations. Script readiness classification can separate safe, manual-review, unsafe, and unmapped controls. Packaging with PyInstaller or another installer tool can make the application easier to distribute for demonstrations. Fleet control may be researched separately only after the local product is complete and a proper multi-machine test lab is available.",
                ],
            ),
        ],
    ),
    (
        "Chapter: 8",
        "REFERENCES",
        [
            (
                "References",
                [
                    "[1] Center for Internet Security, CIS Microsoft Windows 11 Stand-alone Benchmark, Version 4.0.0.",
                    "[2] Center for Internet Security, CIS Ubuntu Linux 24.04 LTS Benchmark, Version 1.0.0.",
                    "[3] National Institute of Standards and Technology, Guide to General Server Security, NIST Special Publication 800-123.",
                    "[4] Python Software Foundation, Python 3 Documentation.",
                    "[5] Qt Company, Qt for Python / PySide6 Documentation.",
                    "[6] SQLite Consortium, SQLite Documentation.",
                    "[7] Ollama Project Documentation, Local Large Language Model Runtime.",
                    "[8] PyMuPDF Documentation, PDF Processing for Python.",
                ],
            ),
        ],
    ),
]


def markdown_text() -> str:
    lines: list[str] = [
        f"# {PROJECT}",
        f"## {SUBTITLE}",
        "",
        "A Project Phase-II Report submitted in partial fulfillment for the award of the degree of Bachelor of Technology in Electronics and Computer Engineering.",
        "",
        "### Students",
    ]
    for name, roll, prn in STUDENTS:
        lines.append(f"- {name} - Roll No: {roll} - PRN No: {prn}")
    lines.extend(
        [
            "",
            f"Supervisor: {SUPERVISOR}",
            f"Department of Electronics Engineering, {COLLEGE}",
            f"Affiliated to {UNIVERSITY}",
            f"Year {YEAR}",
            "",
            "## Abstract",
            "HardSecNet is a local-first cross-platform CIS hardening studio for Windows and Linux. It brings benchmark ingestion, script candidate generation, local profile execution, drift comparison, AI-oriented risk explanation, and evidence-backed reporting into a single PySide6 desktop application.",
            "",
            "## Index",
            "| Ch. No. | Title | Content Description | COs | Page No. |",
            "|---|---|---|---|---|",
        ]
    )
    for row in INDEX_ROWS:
        lines.append("| " + " | ".join(row) + " |")
    for chapter, title, subsections in CHAPTERS:
        lines.append(f"\n# {chapter}: {title}")
        for heading, paras in subsections:
            lines.append(f"\n## {heading}")
            lines.extend(paras)
    for heading, paras in APPENDICES:
        lines.append(f"\n# {heading}")
        lines.extend(paras)
    return "\n\n".join(lines) + "\n"


def add_docx_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    run.bold = bold


def build_docx(path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(PROJECT)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x0B, 0x4F, 0x6C)
    p = doc.add_paragraph(SUBTITLE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    doc.add_paragraph(
        "A Project Phase-II Report submitted in partial fulfillment for the award of the degree of Bachelor of Technology in Electronics and Computer Engineering."
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    doc.add_paragraph("By").alignment = WD_ALIGN_PARAGRAPH.CENTER
    for name, roll, prn in STUDENTS:
        doc.add_paragraph(f"{name}    Roll No: {roll}    PRN No: {prn}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    doc.add_paragraph(f"Under the supervision of {SUPERVISOR}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Department of Electronics Engineering, {COLLEGE}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Affiliated To {UNIVERSITY}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"YEAR {YEAR}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    front = [
        (
            "CERTIFICATE",
            f"This is to certify that the project report entitled '{PROJECT}: {SUBTITLE}' submitted by the listed students to {COLLEGE} is a bonafide record of the project work carried out under the supervision of {SUPERVISOR} during the academic year {YEAR}. The project is approved for the award of the Degree of Bachelor of Technology in Electronics and Computer Engineering.",
        ),
        (
            "DECLARATION",
            "We hereby declare that the work presented in this Project Phase-II report is original and has been carried out by us under the guidance of our project supervisor. This work has not been submitted to any other institute or university for the award of any degree.",
        ),
        (
            "ACKNOWLEDGEMENT",
            f"We express sincere gratitude to {COLLEGE} for providing the opportunity and infrastructure to undertake this project. We are deeply grateful to {SUPERVISOR} for guidance, technical insight, and encouragement. We also thank {HOD}, the department faculty, our peers, and our families for their support.",
        ),
        (
            "ABSTRACT",
            "HardSecNet is a local-first cross-platform CIS hardening studio for Windows and Linux. It brings benchmark ingestion, script candidate generation, local profile execution, drift comparison, AI-oriented risk explanation, and evidence-backed reporting into a single PySide6 desktop application. The final project is software-only and demonstrable on one machine. It intentionally excludes fleet control, remote agents, and remote job orchestration so the implemented scope remains feasible, validated, and aligned with available project resources.",
        ),
    ]
    for heading, body in front:
        h = doc.add_heading(heading, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_docx_paragraph(doc, body)
        if heading == "CERTIFICATE":
            doc.add_paragraph("")
            doc.add_paragraph(f"Supervisor: {SUPERVISOR}")
            doc.add_paragraph(f"Head of Department: {HOD}")
            doc.add_paragraph(f"Principal: {PRINCIPAL}")
        doc.add_page_break()

    h = doc.add_heading("INDEX", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=5)
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = ["Ch. No.", "Title", "Content Description", "COs", "Page No."][i]
    for row in INDEX_ROWS:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_page_break()

    for chapter, title, subsections in CHAPTERS:
        h = doc.add_heading(chapter, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h = doc.add_heading(title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for heading, paras in subsections:
            doc.add_heading(heading, level=2)
            for para in paras:
                add_docx_paragraph(doc, para)
            doc.add_page_break()
        if title.startswith("SYSTEM"):
            add_docx_table(doc, TOOLS_TABLE)
        if title.startswith("TESTING"):
            add_docx_table(doc, TEST_TABLE)
        if title.startswith("PERFORMANCE"):
            add_docx_table(doc, VALIDATION_TABLE)
        doc.add_page_break()
    for heading, paras in APPENDICES:
        h = doc.add_heading(heading, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for para in paras:
            add_docx_paragraph(doc, para)
        doc.add_page_break()
    doc.save(path)


def add_docx_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(rows[0]))
    for i, value in enumerate(rows[0]):
        table.rows[0].cells[i].text = value
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def build_pdf(path: Path) -> None:
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CenterTitle", parent=styles["Title"], alignment=TA_CENTER, fontSize=22, leading=26, textColor=colors.HexColor("#0B4F6C")))
    styles.add(ParagraphStyle(name="CenterSub", parent=styles["Heading2"], alignment=TA_CENTER, fontSize=13, leading=17))
    styles.add(ParagraphStyle(name="Chapter", parent=styles["Heading1"], alignment=TA_CENTER, fontSize=16, leading=21, textColor=colors.HexColor("#0B4F6C")))
    styles.add(ParagraphStyle(name="BodyJustify", parent=styles["BodyText"], alignment=TA_JUSTIFY, fontSize=10.5, leading=15, spaceAfter=8))
    styles.add(ParagraphStyle(name="Small", parent=styles["BodyText"], fontSize=8.6, leading=11))
    story = []

    def p(text: str, style: str = "BodyJustify") -> None:
        story.append(Paragraph(escape(text), styles[style]))

    def h(text: str, level: int = 1) -> None:
        story.append(Paragraph(escape(text), styles["Chapter" if level == 1 else "Heading2"]))

    story.append(Spacer(1, 0.4 * inch))
    story.append(Paragraph(PROJECT, styles["CenterTitle"]))
    story.append(Paragraph(SUBTITLE, styles["CenterSub"]))
    story.append(Spacer(1, 0.2 * inch))
    p("A Project Phase-II Report submitted in partial fulfillment for the award of the degree of Bachelor of Technology in Electronics and Computer Engineering.")
    story.append(Spacer(1, 0.25 * inch))
    story.append(Paragraph("By", styles["CenterSub"]))
    for name, roll, prn in STUDENTS:
        story.append(Paragraph(escape(f"{name}    Roll No: {roll}    PRN No: {prn}"), styles["CenterSub"]))
    story.append(Spacer(1, 0.35 * inch))
    story.append(Paragraph(escape(f"Under the supervision of {SUPERVISOR}"), styles["CenterSub"]))
    story.append(Paragraph(escape(f"Department of Electronics Engineering, {COLLEGE}"), styles["CenterSub"]))
    story.append(Paragraph(escape(f"Affiliated To {UNIVERSITY}"), styles["CenterSub"]))
    story.append(Paragraph(escape(f"YEAR {YEAR}"), styles["CenterSub"]))
    story.append(PageBreak())

    for heading, body in [
        ("CERTIFICATE", f"This is to certify that the project report entitled '{PROJECT}: {SUBTITLE}' submitted by the listed students to {COLLEGE} is a bonafide record of the project work carried out under the supervision of {SUPERVISOR} during the academic year {YEAR}. The project is approved for the award of the Degree of Bachelor of Technology in Electronics and Computer Engineering."),
        ("DECLARATION", "We hereby declare that the work presented in this Project Phase-II report is original and has been carried out by us under the guidance of our project supervisor. This work has not been submitted to any other institute or university for the award of any degree."),
        ("ACKNOWLEDGEMENT", f"We express sincere gratitude to {COLLEGE} for providing the opportunity and infrastructure to undertake this project. We are deeply grateful to {SUPERVISOR} for guidance, technical insight, and encouragement. We also thank {HOD}, the department faculty, our peers, and our families for their support."),
        ("ABSTRACT", "HardSecNet is a local-first cross-platform CIS hardening studio for Windows and Linux. It brings benchmark ingestion, script candidate generation, local profile execution, drift comparison, AI-oriented risk explanation, and evidence-backed reporting into a single PySide6 desktop application. The final project is software-only and demonstrable on one machine. It intentionally excludes fleet control, remote agents, and remote job orchestration so the implemented scope remains feasible, validated, and aligned with available project resources."),
    ]:
        h(heading)
        p(body)
        if heading == "CERTIFICATE":
            story.append(Spacer(1, 0.75 * inch))
            p(f"Supervisor: {SUPERVISOR}     Head of Department: {HOD}     Principal: {PRINCIPAL}")
        story.append(PageBreak())

    h("INDEX")
    table_data = [["Ch. No.", "Title", "Content Description", "COs", "Page No."], *INDEX_ROWS]
    story.append(make_table(table_data, [0.55 * inch, 1.45 * inch, 3.15 * inch, 0.45 * inch, 0.55 * inch], styles))
    story.append(PageBreak())

    for chapter, title, subsections in CHAPTERS:
        h(chapter)
        h(title)
        for heading, paras in subsections:
            h(heading, level=2)
            for para in paras:
                p(para)
            story.append(PageBreak())
        if title.startswith("SYSTEM"):
            story.append(make_table(TOOLS_TABLE, [1.1 * inch, 1.55 * inch, 3.5 * inch], styles))
        if title.startswith("TESTING"):
            story.append(make_table(TEST_TABLE, [0.8 * inch, 4.0 * inch, 1.0 * inch], styles))
        if title.startswith("PERFORMANCE"):
            story.append(make_table(VALIDATION_TABLE, [2.6 * inch, 3.4 * inch], styles))
        story.append(PageBreak())
    for heading, paras in APPENDICES:
        h(heading)
        for para in paras:
            p(para)
        story.append(PageBreak())

    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=0.7 * inch, leftMargin=0.7 * inch, topMargin=0.7 * inch, bottomMargin=0.7 * inch)
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def make_table(rows: list[list[str]], widths: list[float], styles) -> Table:
    table = Table([[Paragraph(escape(cell), styles["Small"]) for cell in row] for row in rows], colWidths=widths)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9EAF2")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ]
        )
    )
    return table


def footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.drawString(0.7 * inch, 0.45 * inch, "HardSecNet BlackBook - Recreated Local-Only Report")
    canvas.drawRightString(A4[0] - 0.7 * inch, 0.45 * inch, str(doc.page))
    canvas.restoreState()


def main() -> None:
    md_path = OUT / "HardSecNet_BlackBook_Recreated_LocalOnly.md"
    docx_path = OUT / "HardSecNet_BlackBook_Recreated_LocalOnly.docx"
    pdf_path = OUT / "HardSecNet_BlackBook_Recreated_LocalOnly.pdf"
    md_path.write_text(markdown_text(), encoding="utf-8")
    build_docx(docx_path)
    build_pdf(pdf_path)
    print(md_path)
    print(docx_path)
    print(pdf_path)


if __name__ == "__main__":
    main()
